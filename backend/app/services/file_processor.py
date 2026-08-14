"""
File Processor
Extracts text content from uploaded files for LLM context injection.
No chunking, no embedding, no vector DB — just text extraction.
"""
import logging
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

_ocr_engine_instance = None


def _get_ocr_engine():
    """Lazy, cached RapidOCR engine (avoid reloading the model per file)."""
    global _ocr_engine_instance
    if _ocr_engine_instance is None:
        from rapidocr import RapidOCR
        _ocr_engine_instance = RapidOCR()
    return _ocr_engine_instance


class FileProcessor:
    """Extracts text content from files."""

    def process_file(self, file_path: str, conversation_id: str, user_id: int = 1) -> Dict[str, Any]:
        """Extract text from a file and store reference."""
        path = Path(file_path).resolve()
        suffix = path.suffix.lower()

        try:
            text = self._extract_text(path, suffix)
            self._store_media_text(file_path, text)
            return {"status": "success", "file_name": path.name, "text_length": len(text), "text": text}
        except Exception as e:
            logger.error(f"File processing failed: {e}")
            return {"status": "error", "message": str(e)}

    def _extract_text(self, path: Path, suffix: str) -> str:
        """Extract text based on file type."""
        # Plain text
        if suffix in ('.txt', '.md', '.py', '.js', '.ts', '.html', '.css', '.json', '.csv', '.xml', '.yaml', '.yml'):
            return path.read_text(errors='ignore')

        # PDF
        if suffix == '.pdf':
            return self._extract_pdf(path)

        # Word
        if suffix == '.docx':
            return self._extract_docx(path)

        # Images — OCR text (if OCR installed), otherwise note for vision
        if suffix in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'):
            try:
                from PIL import Image
                pil = Image.open(str(path))
                text = self._ocr_image(pil)
                return text.strip() if text and text.strip() else "[Image — no text detected; may need vision analysis]"
            except ImportError:
                return "[Image — OCR (rapidocr/pillow) not installed]"

        # Spreadsheets — read cell values into a tabular text form (openpyxl/xlrd).
        if suffix in ('.xlsx', '.xlsm'):
            return self._extract_xlsx(path)
        if suffix == '.xls':
            return self._extract_xls(path)

        # Audio — no text extraction here; ASR transcription happens in the
        # multimodal ingestion path (Architecture B).
        if suffix in ('.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac'):
            return "[Audio — transcribed via ASR in multimodal ingestion]"

        return "[Unsupported file format]"

    def _extract_xlsx(self, path: Path) -> str:
        """Read .xlsx cell values into a TSV-like text (openpyxl)."""
        try:
            import openpyxl
        except ImportError:
            return "[Spreadsheet — openpyxl not installed]"
        try:
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                lines.append(f"--- Sheet: {ws.title} ---")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append("\t".join(cells))
            wb.close()
            return "\n".join(lines)
        except Exception as exc:
            logger.warning("xlsx extraction failed: %s", exc)
            return "[Spreadsheet — failed to read]"

    def _extract_xls(self, path: Path) -> str:
        """Read legacy .xls cell values (xlrd)."""
        try:
            import xlrd
        except ImportError:
            return "[Spreadsheet — xlrd not installed]"
        try:
            book = xlrd.open_workbook(str(path))
            lines: list[str] = []
            for sheet in book.sheets():
                lines.append(f"--- Sheet: {sheet.name} ---")
                for r in range(sheet.nrows):
                    cells = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                    cells = [c for c in cells if c.strip()]
                    if cells:
                        lines.append("\t".join(cells))
            return "\n".join(lines)
        except Exception as exc:
            logger.warning("xls extraction failed: %s", exc)
            return "[Spreadsheet — failed to read]"

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF. If the PDF is scanned (no embedded text),
        fall back to rendering pages and OCR-ing them (pypdfium2 + RapidOCR)."""
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
                text = "\n\n".join(pages)
        except ImportError:
            logger.warning("pdfplumber not installed, trying pypdf")
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                text = "\n\n".join(p.extract_text() for p in reader.pages if p.extract_text())
            except ImportError:
                return "[PDF — install pypdf or pdfplumber for text extraction]"
        except Exception as exc:
            logger.warning("PDF text extraction failed: %s", exc)

        if text and text.strip():
            return text

        # Scanned PDF (no embedded text) — OCR the rendered pages.
        logger.info("No embedded text in %s — running OCR", path.name)
        return self._ocr_pdf(path)

    def _ocr_pdf(self, path: Path) -> str:
        """Render each PDF page to an image (pypdfium2) and OCR it (RapidOCR)."""
        try:
            import pypdfium2 as pdfium
        except ImportError:
            return "[PDF — install pypdfium2 for scanned-page OCR]"
        try:
            pdf = pdfium.PdfDocument(str(path))
        except Exception as exc:
            logger.warning("Could not open PDF for OCR: %s", exc)
            return ""
        parts = []
        try:
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=2.0)  # ~144 DPI
                try:
                    pil = bitmap.to_pil().convert("RGB")
                finally:
                    bitmap.close()
                page_text = self._ocr_image(pil)
                if page_text and page_text.strip():
                    parts.append(page_text)
        except Exception as exc:
            logger.warning("PDF OCR failed: %s", exc)
        finally:
            pdf.close()
        return "\n\n".join(parts)

    def _ocr_image(self, pil_image) -> str:
        """OCR a PIL image with RapidOCR (cached engine). Returns '' on failure."""
        try:
            from rapidocr import RapidOCR
        except ImportError:
            logger.warning("RapidOCR not installed")
            return ""
        try:
            engine = _get_ocr_engine()
            out = engine(pil_image)
            if out is None:
                return ""
            texts = getattr(out, "txts", None) or []
            return "\n".join(str(t).strip() for t in texts if t and str(t).strip())
        except Exception as exc:
            logger.warning("OCR failed: %s", exc)
            return ""

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return "[DOCX — install python-docx for text extraction]"

    def _store_media_text(self, file_path: str, text: str) -> None:
        """Store extracted text in media_library."""
        from app.database import get_db
        db = get_db()
        db.execute(
            "UPDATE media_library SET extracted_text = ? WHERE file_path = ?",
            (text[:50000], file_path)  # cap at 50K chars
        )
        db.commit()
